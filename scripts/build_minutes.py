# -*- coding: utf-8 -*-
"""
build_minutes.py · 基于 assets/templates/访谈纪要输出格式.docx 骨架生成纪要 docx + md。

调用方式：
    from build_minutes import build_minutes
    build_minutes(
        company="星辰新能",
        meta={"访谈对象": "...", "时间": "...", "地点": "...", "人员": "..."},
        rows=[...20 条字符串...],       # 顺序参见 SKILL.md 第 6 节
        key_points=[...],               # 可选，3–6 条
        followups=[...],                # 可选，每条带责任方/时点
        out_dir="/path/to/output",
    )

每条 rows 内可以用 "\n" 分段；row[14]、row[18] 期望前段写本主题、空行后用 "[各行业的收入占比] xxx" / "[业绩对赌] xxx" 写并列子项。
"""
from pathlib import Path
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

TEMPLATE = Path(__file__).resolve().parent.parent / "assets" / "templates" / "访谈纪要输出格式.docx"

DIM_LABELS = [
    ("1. 行业", "主营业务"),
    ("1. 行业", "市场痛点 & 市场规模"),
    ("1. 行业", "竞争格局及发展趋势"),
    ("1. 行业", "公司竞争优劣势"),
    ("2. 团队", "实控人履历和创业背景"),
    ("2. 团队", "管理层其他人员情况"),
    ("2. 团队", "员工构成"),
    ("2. 团队", "股权结构"),
    ("2. 团队", "公司战略目标"),
    ("3. 产品业务", "业务构成 & 市场策略"),
    ("3. 产品业务", "技术原理及发展趋势"),
    ("3. 产品业务", "产品研发规划"),
    ("3. 产品业务", "产能情况"),
    ("3. 产品业务", "供应链 & 客户情况"),
    ("4. 财务及融资", "历史收入及收入预测 / 各行业的收入占比"),
    ("4. 财务及融资", "毛利率 & 净利率水平"),
    ("4. 财务及融资", "上轮融资情况"),
    ("4. 财务及融资", "本轮融资计划"),
    ("4. 财务及融资", "上市规划 & 中介机构 / 业绩对赌"),
    ("5.", "主要风险"),
]


def _set_run_font(run, font_name: str, size_pt: float, bold: bool = False):
    """对单个 run 设置字体、字号、加粗（中文三个命名空间统一设）。"""
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)


def _set_cell_text(cell, content: str):
    """清空 cell，按换行分段写入，表格内容：黑体 五号 (10.5pt)。"""
    paras = list(cell.paragraphs)
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    first = cell.paragraphs[0]
    for r in list(first.runs):
        r._element.getparent().remove(r._element)
    lines = content.split("\n")
    run = first.add_run(lines[0])
    _set_run_font(run, '黑体', 10.5)
    for line in lines[1:]:
        p = cell.add_paragraph(line)
        for r in p.runs:
            _set_run_font(r, '黑体', 10.5)


def _add_section_heading(doc, text: str):
    """章节标题：黑体 14pt 加粗。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, '黑体', 14, bold=True)
    return p


def build_minutes(
    company: str,
    meta: Dict[str, str],
    rows: List[str],
    key_points: Optional[List[str]] = None,
    followups: Optional[List[str]] = None,
    out_dir: str = ".",
) -> Dict[str, str]:
    assert len(rows) == 20, f"rows must be 20 items, got {len(rows)}"
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    # ---- docx ----
    doc = Document(str(TEMPLATE))

    # 标题：黑体 18pt 加粗
    title_p = doc.paragraphs[0]
    title_p.text = f"{company}访谈纪要"
    for run in title_p.runs:
        _set_run_font(run, '黑体', 18, bold=True)

    # 元信息行：黑体 12pt
    for i, k in enumerate(["访谈对象", "时间", "地点", "人员"], start=1):
        meta_p = doc.paragraphs[i]
        meta_p.text = f"{k}：{meta.get(k, '')}"
        for run in meta_p.runs:
            _set_run_font(run, '黑体', 12)

    table = doc.tables[0]
    assert len(table.rows) == 20
    for r_idx, content in enumerate(rows):
        _set_cell_text(table.rows[r_idx].cells[2], content)

    if key_points:
        doc.add_paragraph("")
        _add_section_heading(doc, "核心观点摘要")
        for i, kp in enumerate(key_points, 1):
            p = doc.add_paragraph(f"{i}. {kp}")
            for run in p.runs:
                _set_run_font(run, '黑体', 12)

    if followups:
        doc.add_paragraph("")
        _add_section_heading(doc, "后续事项")
        for fu in followups:
            p = doc.add_paragraph(f"- {fu}")
            for run in p.runs:
                _set_run_font(run, '黑体', 12)

    out_docx = out_dir / f"{company}访谈纪要.docx"
    doc.save(str(out_docx))

    # ---- md ----
    md = []
    md.append(f"# {company}访谈纪要\n")
    for k in ["访谈对象", "时间", "地点", "人员"]:
        md.append(f"**{k}**：{meta.get(k, '')}  ")
    md.append("\n---\n")
    md.append("## 一、访谈内容\n")
    md.append("| 类别 | 考察维度 | 内容 |")
    md.append("|---|---|---|")
    for (cat, dim), content in zip(DIM_LABELS, rows):
        md.append(f"| {cat} | {dim} | {content.replace(chr(10), '<br>')} |")
    if key_points:
        md.append("\n---\n")
        md.append("## 二、核心观点摘要\n")
        for i, kp in enumerate(key_points, 1):
            md.append(f"{i}. {kp}")
    if followups:
        md.append("\n---\n")
        md.append("## 三、后续事项\n")
        for fu in followups:
            md.append(f"- {fu}")

    out_md = out_dir / f"{company}访谈纪要.md"
    out_md.write_text("\n".join(md), encoding="utf-8")

    return {"docx": str(out_docx), "md": str(out_md)}


if __name__ == "__main__":
    print("This is a helper module. Import build_minutes() from your skill workflow.")
    print(f"Template at: {TEMPLATE}")
